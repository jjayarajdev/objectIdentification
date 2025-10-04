"""
Report Generation Service
Creates professional Word documents and PDFs from room analysis
"""

from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import io
import base64
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image


class ReportGenerator:
    """Service for generating professional reports from room analysis"""

    def __init__(self):
        self.company_name = "Property & Scene Analysis System"
        self.company_tagline = "AI-Powered Surveyor Intelligence Platform"

    def create_word_report(self, analysis: Dict[str, Any], image_path: Path = None) -> io.BytesIO:
        """
        Create a professional Word document report

        Args:
            analysis: Room analysis data
            image_path: Path to the analyzed image

        Returns:
            BytesIO object containing the Word document
        """
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Room Intelligence Analysis Report"
        doc.core_properties.subject = "Comprehensive Room Analysis"
        doc.core_properties.author = self.company_name
        doc.core_properties.comments = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

        # Add custom styles
        self._setup_styles(doc)

        # Add header
        self._add_header(doc)

        # Add executive summary section
        self._add_executive_summary(doc, analysis)

        # Add image if provided
        if image_path and image_path.exists():
            self._add_image_section(doc, image_path)

        # Add people analysis
        self._add_people_section(doc, analysis.get('people_analysis', {}))

        # Add environmental conditions
        self._add_environment_section(doc, analysis.get('temperature', {}),
                                    analysis.get('hvac_and_blinds', {}))

        # Add furniture inventory
        self._add_furniture_section(doc, analysis.get('furniture', []))

        # Add lighting analysis
        self._add_lighting_section(doc, analysis.get('lighting', []))

        # Add flooring details
        self._add_flooring_section(doc, analysis.get('flooring', {}))

        # Add electronics inventory
        self._add_electronics_section(doc, analysis.get('electronics', []))

        # Add plants section
        self._add_plants_section(doc, analysis.get('plants', []))

        # Add room metrics
        self._add_metrics_section(doc, analysis.get('room_metrics', {}))

        # Add cost summary with charts
        self._add_cost_summary_section(doc, analysis.get('cost_summary', {}))

        # Add room classification
        self._add_classification_section(doc, analysis.get('room_classification', {}))

        # Add detailed narrative
        self._add_narrative_section(doc, analysis.get('detailed_narrative', ''))

        # Add footer
        self._add_footer(doc)

        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io

    def _setup_styles(self, doc):
        """Setup custom styles for the document"""
        styles = doc.styles

        # Heading 1 style
        heading1 = styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(18)
        heading1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        heading1.font.bold = True

        # Heading 2 style
        heading2 = styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(14)
        heading2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        heading2.font.bold = True

        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

    def _add_header(self, doc):
        """Add document header"""
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = self.company_name
        header_para.style.font.size = Pt(14)
        header_para.style.font.bold = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add title
        title = doc.add_heading('Room Intelligence Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add tagline
        tagline = doc.add_paragraph(self.company_tagline)
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline.style.font.italic = True

        # Add generation date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _add_executive_summary(self, doc, analysis):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', 1)

        # Create summary bullets
        summary_points = []

        if 'people_analysis' in analysis:
            summary_points.append(f"👥 {analysis['people_analysis']['total_count']} people present in the room")

        if 'room_classification' in analysis:
            summary_points.append(f"🏢 Room Type: {analysis['room_classification']['primary_use']}")
            summary_points.append(f"🎨 Design Style: {analysis['room_classification']['design_style']}")

        if 'cost_summary' in analysis:
            min_cost = analysis['cost_summary']['complete_room_estimate']['min']
            max_cost = analysis['cost_summary']['complete_room_estimate']['max']
            summary_points.append(f"💰 Estimated Total Value: ₹{min_cost:,.0f} - ₹{max_cost:,.0f}")

        if 'room_metrics' in analysis:
            area = analysis['room_metrics']['total_area_sqft']
            summary_points.append(f"📐 Total Area: {area} sq ft")

        for point in summary_points:
            doc.add_paragraph(point, style='List Bullet')

        doc.add_paragraph()

    def _add_image_section(self, doc, image_path):
        """Add the analyzed image"""
        doc.add_heading('Analyzed Image', 1)
        doc.add_picture(str(image_path), width=Inches(6))
        doc.add_paragraph()

    def _add_people_section(self, doc, people_data):
        """Add people analysis section"""
        doc.add_heading('🧍 People Analysis', 1)

        if people_data:
            doc.add_paragraph(f"Total Count: {people_data.get('total_count', 0)} people")
            doc.add_paragraph(f"• Sitting: {people_data.get('sitting', 0)}")
            doc.add_paragraph(f"• Standing: {people_data.get('standing', 0)}")

            if 'details' in people_data:
                doc.add_paragraph()
                doc.add_paragraph(people_data['details'])

        doc.add_paragraph()

    def _add_environment_section(self, doc, temp_data, hvac_data):
        """Add environmental conditions section"""
        doc.add_heading('🌡️ Environmental Conditions', 1)

        if temp_data:
            doc.add_paragraph('Temperature Analysis', style='Heading 2')
            temp_range = temp_data.get('estimated_range_celsius', {})
            doc.add_paragraph(f"• Estimated Range: {temp_range.get('min', 'N/A')}°C - {temp_range.get('max', 'N/A')}°C")
            doc.add_paragraph(f"• HVAC Setting: {temp_data.get('hvac_setting', 'N/A')}")
            doc.add_paragraph(f"• Comfort Assessment: {temp_data.get('comfort_assessment', 'N/A')}")

        if hvac_data:
            doc.add_paragraph('HVAC & Window Treatments', style='Heading 2')
            doc.add_paragraph(f"• AC Type: {hvac_data.get('ac_type', 'N/A')}")
            doc.add_paragraph(f"• Brand: {hvac_data.get('brand', 'Unknown')}")
            doc.add_paragraph(f"• Blinds Type: {hvac_data.get('blinds_type', 'N/A')}")
            doc.add_paragraph(f"• Blinds Material: {hvac_data.get('blinds_material', 'N/A')}")

        doc.add_paragraph()

    def _add_furniture_section(self, doc, furniture_list):
        """Add furniture inventory section"""
        doc.add_heading('🪑 Furniture Inventory', 1)

        if furniture_list:
            # Create table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Add headers
            headers = ['Item', 'Quantity', 'Material', 'Condition', 'Estimated Cost (₹)']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Add data
            for item in furniture_list:
                row = table.add_row().cells
                row[0].text = item.get('item', '')
                row[1].text = str(item.get('quantity', 1))
                row[2].text = item.get('material', '')
                row[3].text = item.get('condition', '')
                cost = item.get('estimated_cost_inr', {})
                row[4].text = f"{cost.get('min', 0):,} - {cost.get('max', 0):,}"

        doc.add_paragraph()

    def _add_lighting_section(self, doc, lighting_list):
        """Add lighting analysis section"""
        doc.add_heading('💡 Lighting Systems', 1)

        if lighting_list:
            for light in lighting_list:
                doc.add_paragraph(f"• {light.get('type', 'Unknown')} ({light.get('quantity', 1)} units)")
                doc.add_paragraph(f"  Style: {light.get('style', 'N/A')}", style='List Continue')
                cost = light.get('estimated_cost_inr', {})
                doc.add_paragraph(f"  Cost: ₹{cost.get('min', 0):,} - ₹{cost.get('max', 0):,}", style='List Continue')

        doc.add_paragraph()

    def _add_flooring_section(self, doc, flooring_data):
        """Add flooring details section"""
        doc.add_heading('🏢 Flooring Details', 1)

        if flooring_data:
            doc.add_paragraph(f"• Type: {flooring_data.get('type', 'N/A')}")
            doc.add_paragraph(f"• Material: {flooring_data.get('material', 'N/A')}")
            doc.add_paragraph(f"• Pattern: {flooring_data.get('pattern', 'N/A')}")
            cost = flooring_data.get('cost_per_sqft_inr', {})
            doc.add_paragraph(f"• Cost per sq ft: ₹{cost.get('min', 0)} - ₹{cost.get('max', 0)}")
            doc.add_paragraph(f"• Total Area: {flooring_data.get('total_area_sqft', 0)} sq ft")

        doc.add_paragraph()

    def _add_electronics_section(self, doc, electronics_list):
        """Add electronics inventory section"""
        doc.add_heading('📱 Electronics Inventory', 1)

        if electronics_list:
            for device in electronics_list:
                doc.add_paragraph(f"• {device.get('device', 'Unknown')} ({device.get('quantity', 1)} units)")
                doc.add_paragraph(f"  Brand (estimated): {device.get('brand_guess', 'Unknown')}", style='List Continue')
                cost = device.get('estimated_cost_inr', {})
                doc.add_paragraph(f"  Value: ₹{cost.get('min', 0):,} - ₹{cost.get('max', 0):,}", style='List Continue')

        doc.add_paragraph()

    def _add_plants_section(self, doc, plants_list):
        """Add plants section"""
        doc.add_heading('🌿 Plants', 1)

        if plants_list:
            for plant in plants_list:
                doc.add_paragraph(f"• {plant.get('common_name', 'Unknown')}")
                doc.add_paragraph(f"  Scientific: {plant.get('scientific_name', 'N/A')}", style='List Continue')
                doc.add_paragraph(f"  Size: {plant.get('size', 'N/A')}", style='List Continue')
                cost = plant.get('estimated_cost_inr', {})
                doc.add_paragraph(f"  Cost: ₹{cost.get('min', 0):,} - ₹{cost.get('max', 0):,}", style='List Continue')

        doc.add_paragraph()

    def _add_metrics_section(self, doc, metrics_data):
        """Add room metrics section"""
        doc.add_heading('📐 Room Metrics', 1)

        if metrics_data:
            dimensions = metrics_data.get('estimated_dimensions_feet', {})
            doc.add_paragraph(f"• Dimensions: {dimensions.get('length', 0)}' × {dimensions.get('width', 0)}'")
            doc.add_paragraph(f"• Ceiling Height: {metrics_data.get('ceiling_height_feet', 0)}'")
            doc.add_paragraph(f"• Total Area: {metrics_data.get('total_area_sqft', 0)} sq ft")

        doc.add_paragraph()

    def _add_cost_summary_section(self, doc, cost_data):
        """Add cost summary section with table"""
        doc.add_heading('💰 Cost Summary', 1)

        if cost_data:
            # Create table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Medium Shading 1 Accent 1'

            # Add headers
            headers = ['Category', 'Minimum (₹)', 'Maximum (₹)']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Add data rows
            categories = [
                ('Furniture', 'furniture_total'),
                ('Lighting', 'lighting_total'),
                ('Flooring', 'flooring_total'),
                ('Plants', 'plants_total'),
                ('TOTAL ESTIMATE', 'complete_room_estimate')
            ]

            for cat_name, cat_key in categories:
                if cat_key in cost_data:
                    row = table.add_row().cells
                    row[0].text = cat_name
                    row[1].text = f"₹{cost_data[cat_key]['min']:,.0f}"
                    row[2].text = f"₹{cost_data[cat_key]['max']:,.0f}"

                    # Bold the total row
                    if cat_name == 'TOTAL ESTIMATE':
                        for cell in row:
                            cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

    def _add_classification_section(self, doc, classification_data):
        """Add room classification section"""
        doc.add_heading('🏢 Room Classification', 1)

        if classification_data:
            doc.add_paragraph(f"• Primary Use: {classification_data.get('primary_use', 'N/A')}")
            doc.add_paragraph(f"• Design Style: {classification_data.get('design_style', 'N/A')}")
            doc.add_paragraph(f"• Formality Level: {classification_data.get('formality_level', 'N/A')}")
            doc.add_paragraph(f"• Organization Type: {classification_data.get('organization_type', 'N/A')}")

        doc.add_paragraph()

    def _add_narrative_section(self, doc, narrative):
        """Add detailed narrative section"""
        doc.add_heading('📝 Detailed Analysis', 1)

        if narrative:
            doc.add_paragraph(narrative)

        doc.add_paragraph()

    def _add_footer(self, doc):
        """Add document footer"""
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"© {datetime.now().year} {self.company_name} | Confidential Report"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.style.font.size = Pt(9)
        footer_para.style.font.italic = True

    def create_cost_chart(self, cost_data):
        """Create a cost breakdown chart"""
        categories = []
        min_costs = []
        max_costs = []

        for category, key in [
            ('Furniture', 'furniture_total'),
            ('Lighting', 'lighting_total'),
            ('Flooring', 'flooring_total'),
            ('Plants', 'plants_total')
        ]:
            if key in cost_data:
                categories.append(category)
                min_costs.append(cost_data[key]['min'])
                max_costs.append(cost_data[key]['max'])

        if not categories:
            return None

        # Create bar chart
        fig, ax = plt.subplots(figsize=(10, 6))
        x = np.arange(len(categories))
        width = 0.35

        bars1 = ax.bar(x - width/2, min_costs, width, label='Minimum', color='#2E74B5')
        bars2 = ax.bar(x + width/2, max_costs, width, label='Maximum', color='#70AD47')

        ax.set_xlabel('Category')
        ax.set_ylabel('Cost (₹)')
        ax.set_title('Cost Breakdown by Category')
        ax.set_xticks(x)
        ax.set_xticklabels(categories)
        ax.legend()

        # Add value labels on bars
        for bars in [bars1, bars2]:
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'₹{height:,.0f}',
                       ha='center', va='bottom', fontsize=9)

        plt.tight_layout()

        # Save to BytesIO
        img_io = io.BytesIO()
        plt.savefig(img_io, format='png', dpi=150)
        img_io.seek(0)
        plt.close()

        return img_io

    def create_scene_analysis_report(self, analysis: Dict[str, Any], image_path: Path = None) -> io.BytesIO:
        """
        Create a Word document for scene analysis (surveyor images)

        Args:
            analysis: Scene analysis data with narrative_report
            image_path: Optional path to analyzed image

        Returns:
            BytesIO containing Word document
        """
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Property & Scene Analysis Report"
        doc.core_properties.author = self.company_name
        doc.core_properties.comments = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

        # Setup styles
        self._setup_styles(doc)

        # Add title
        title = doc.add_heading('Property & Scene Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add company tagline
        tagline = doc.add_paragraph(self.company_tagline)
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline.style.font.italic = True

        # Add generation date
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add scene type
        scene_type_label = self._get_scene_type_label(analysis.get('scene_type', 'Unknown'))
        scene_para = doc.add_paragraph()
        scene_para.add_run('Scene Type: ').bold = True
        scene_para.add_run(scene_type_label)
        scene_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Add spacing

        # Add scene overview if present
        if analysis.get('scene_overview'):
            doc.add_heading('Scene Overview', 1)
            doc.add_paragraph(analysis['scene_overview'])
            doc.add_paragraph()

        # Add narrative report if present (the detailed analysis)
        if analysis.get('narrative_report'):
            doc.add_heading('Detailed Analysis', 1)
            self._add_narrative_markdown(doc, analysis['narrative_report'])
            doc.add_paragraph()

        # Add simplified data table if present
        if analysis.get('simplified_data'):
            doc.add_heading('Summary Table', 1)
            self._add_simplified_table(doc, analysis['simplified_data'])
            doc.add_paragraph()

        # Add key observations
        if analysis.get('key_observations') and len(analysis['key_observations']) > 0:
            doc.add_heading('Key Observations', 1)
            for obs in analysis['key_observations']:
                doc.add_paragraph(f"• {obs}", style='List Bullet')
            doc.add_paragraph()

        # Add property value estimate
        if analysis.get('estimated_property_value'):
            value_data = analysis['estimated_property_value']
            if value_data.get('min') or value_data.get('max'):
                doc.add_heading('Estimated Property Value', 1)
                p = doc.add_paragraph()
                min_val = value_data.get('min', 0)
                max_val = value_data.get('max', 0)
                p.add_run(f'₹{min_val:,.0f} - ₹{max_val:,.0f}').bold = True
                p.add_run(f"\n{value_data.get('basis', '')}")
                doc.add_paragraph()

        # Add footer
        self._add_footer(doc)

        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return doc_io

    def _get_scene_type_label(self, scene_type: str) -> str:
        """Get friendly label for scene type"""
        labels = {
            'indoor_office': 'Indoor Office',
            'indoor_industrial': 'Industrial Space',
            'building_exterior': 'Building Exterior',
            'land_property': 'Land/Property',
            'construction_site': 'Construction Site',
            'infrastructure': 'Infrastructure',
            'agricultural': 'Agricultural',
            'natural_landscape': 'Natural Landscape',
            'parking_area': 'Parking Area'
        }
        return labels.get(scene_type, 'Other')

    def _add_narrative_markdown(self, doc, narrative: str):
        """Convert narrative markdown to Word format"""
        lines = narrative.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # Handle emoji headers with bold text
            if any(line.startswith(emoji) for emoji in ['🧍', '🌡️', '🪑', '💡', '🌬️', '💰', '🧶', '📱', '🌿', '✨', '📐', '🏢']):
                p = doc.add_paragraph()
                # Split by ** to handle bold sections
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True
            # Handle bullet points
            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                text = line[1:].strip()
                doc.add_paragraph(text, style='List Bullet')
            # Handle numbered lists
            elif line and line[0].isdigit() and '.' in line[:3]:
                text = line.split('.', 1)[1].strip() if '.' in line else line
                doc.add_paragraph(text, style='List Number')
            # Regular text
            else:
                doc.add_paragraph(line)

    def _add_simplified_table(self, doc, data: list):
        """Add simplified analysis table"""
        if not data:
            return

        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Add headers
        headers = ['Identifier', 'Details', 'Estimated Cost']
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        # Add data rows
        for item in data:
            row = table.add_row().cells
            row[0].text = item.get('identifier', '')
            row[1].text = item.get('details', '')
            row[2].text = item.get('estimated_cost', '—')